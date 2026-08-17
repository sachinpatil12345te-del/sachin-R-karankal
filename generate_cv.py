from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent
IMAGE_PATH = ROOT / "ae1c50ea0808b5c9713262096ec880877b22d108595d8da593c5ed6d5f5ffd43.png"
OUTPUT_PATH = ROOT / "Sachin_R_Karankal_CV.docx"


def add_section_heading(document: Document, title: str):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(13)
    paragraph.space_before = Pt(10)
    paragraph.space_after = Pt(4)


doc = Document()
section = doc.sections[0]
section.left_margin = Inches(0.7)
section.right_margin = Inches(0.7)
section.top_margin = Inches(0.6)
section.bottom_margin = Inches(0.6)

# Header row with photo and personal information
header_table = doc.add_table(rows=1, cols=2)
header_table.style = "Table Grid"
left_cell, right_cell = header_table.rows[0].cells
left_cell.width = Inches(1.9)
right_cell.width = Inches(5.4)

left_paragraph = left_cell.paragraphs[0]
left_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
if IMAGE_PATH.exists():
    photo_run = left_paragraph.add_run()
    photo_run.add_picture(str(IMAGE_PATH), width=Inches(1.5))

right_paragraph = right_cell.paragraphs[0]
name_run = right_paragraph.add_run("Sachin R Karankal")
name_run.bold = True
name_run.font.size = Pt(24)
right_cell.add_paragraph("Career-focused and results-driven professional")
right_cell.add_paragraph("Mobile: +91 9226313805")
right_cell.add_paragraph("WhatsApp: +91 9226313805")
right_cell.add_paragraph("Instagram: @patil_sachin_sg")
right_cell.add_paragraph("Email: sachinpatil12345te@gmail.com")

add_section_heading(doc, "Professional Summary")
doc.add_paragraph(
    "Dedicated and professional individual with strong communication skills, a passion for growth, and a commitment to delivering quality work. "
    "I have practical exposure to AI & ML, Power BI, and data-driven thinking, and I value consistency, responsibility, and continuous learning in every opportunity I take."
)

add_section_heading(doc, "Education")
doc.add_paragraph("Diploma in Computer Engineering")
doc.add_paragraph("R C Patel College of Engineering and Polytechnic, Shirpur")
doc.add_paragraph("Currently in 3rd year of study")
doc.add_paragraph("Focused on hands-on technical learning and practical industry exposure")

add_section_heading(doc, "Skills")
for skill in [
    "Communication",
    "Leadership",
    "Problem Solving",
    "Time Management",
    "Teamwork",
    "Customer Support",
    "Adaptability",
    "Professionalism",
    "AI & ML",
    "Power BI",
    "Python",
    "Data Analysis",
]:
    doc.add_paragraph(f"- {skill}")

add_section_heading(doc, "Internship Experience")
doc.add_paragraph("Internship in AI & ML")
doc.add_paragraph(
    "Completed internship training in Artificial Intelligence and Machine Learning, gaining practical exposure to modern technology, model concepts, and data-driven problem solving."
)

add_section_heading(doc, "Projects")
projects = [
    (
        "Fashion Intelligence Hub",
        "Developed a style prediction and best marketplace recommendation system that helps users discover fashion trends and suitable products based on their preferences.",
    ),
    (
        "Credit Card Fraud Detection Dashboard",
        "Built a dashboard that predicts credit card fraud and displays all relevant transaction and fraud-related information for analysis and decision-making.",
    ),
]
for title, description in projects:
    doc.add_paragraph(title)
    doc.add_paragraph(description)

# Slight professional finishing touch
for paragraph in doc.paragraphs:
    if paragraph.text.strip() and paragraph.text.strip() not in {"", "-"}:
        if not paragraph.runs:
            continue
        for run in paragraph.runs:
            if paragraph.text.strip().startswith("-"):
                run.font.size = Pt(11)
            else:
                run.font.size = Pt(11)

# Save final CV

doc.save(OUTPUT_PATH)
print(f"Professional CV saved to: {OUTPUT_PATH}")
